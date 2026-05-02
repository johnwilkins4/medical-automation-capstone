"""Build the capstone report as a Word document.

Run from the project root:
    python report/build_report.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).parent / "capstone_report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)


def add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)


def add_code(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    # Set a grey-ish background shading on the paragraph.
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)


def add_table(doc, headers: list[str], rows: list[list[str]],
              col_widths_in: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        _set_cell_shading(hdr_cells[i], "D9E8F5")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)

    if col_widths_in:
        for row in table.rows:
            for j, w in enumerate(col_widths_in):
                row.cells[j].width = Inches(w)


# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------

def build() -> Path:
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tighten heading colors
    for h_name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[h_name]
        s.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        s.font.name = "Calibri"

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Agentic Healthcare Assistant\nfor Medical Task Automation")
    r.bold = True
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Applied Generative AI Specialisation — Capstone Report")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"John Wilkins  |  {date.today().isoformat()}\n"
                 "Purdue University / Simplilearn").font.size = Pt(11)

    doc.add_paragraph()

    # ---- 1. Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "This capstone delivers an Agentic Healthcare Assistant that autonomously handles four "
        "interdependent medical-administration tasks — appointment booking, patient-record "
        "management, medical-history retrieval, and trusted medical information search — behind "
        "a single conversational interface. The system is built with a LangGraph planner-executor "
        "agent, OpenAI GPT-4o-mini, a FAISS vector store over both patient history notes and a "
        "seeded Medline/WHO corpus, and a Streamlit dashboard that exposes agent traces and "
        "evaluation metrics. The flagship demo scenario — a caregiver booking a nephrologist for a "
        "70-year-old father with chronic kidney disease while requesting a treatment summary — is "
        "executed end-to-end using all four capabilities in a single user turn.")
    add_para(doc,
        "A purpose-built authorization layer protects patient health information (PHI). Every "
        "PHI-touching tool — record retrieval, record update, and appointment booking — passes "
        "through a code-level hard gate that requires the caller to be on the patient's "
        "authorized-contacts list. A second hard gate prevents the agent from leaking PHI by "
        "inference: the medical-info search refuses to research a condition the user did not "
        "themselves name, blocking the subtle pattern where the agent reads a patient's record, "
        "infers their diagnosis, and 'helpfully' surfaces it as general health education. "
        "Verification persists across follow-up turns within a session, so a verified caregiver "
        "is not re-prompted for credentials between related questions.")

    # ---- 2. Problem Statement ----
    add_heading(doc, "2. Problem Statement", level=1)
    add_para(doc,
        "Modern digital health ecosystems rely on siloed tools for scheduling, recordkeeping, and "
        "clinical lookup, producing a fragmented experience for patients and their caregivers. "
        "Agentic AI — large language models combined with Retrieval-Augmented Generation (RAG) and "
        "memory — offers the potential to coordinate these tasks behind a single interface.")
    add_para(doc, "The capstone requires a system that can:")
    for b in [
        "Book medical appointments by matching patient intent with doctor availability.",
        "Manage medical records — create and update structured or unstructured patient histories.",
        "Retrieve and summarize past diagnoses, treatments, and alerts.",
        "Search trusted external sources (Medline, WHO) and synthesize cited answers.",
    ]:
        add_bullet(doc, b)
    add_para(doc,
        "The flagship scenario given in the problem statement is:", italic=True)
    add_code(doc,
        '"My 70-year-old father has chronic kidney disease. I want to book a nephrologist\n'
        ' for him. Also, can you summarize latest treatment methods?"')
    add_para(doc,
        "This single request exercises all four capabilities and is the system's primary "
        "evaluation target.")

    # ---- 3. System Architecture ----
    add_heading(doc, "3. System Architecture", level=1)
    add_para(doc,
        "The system is organized as a three-node LangGraph: a planner that decomposes the user "
        "query, an executor that runs each tool in sequence, and a composer that writes the final "
        "patient-friendly reply. State flowing through the graph carries the user query, the "
        "resolved patient id, the planner's ordered plan, per-step tool results, a structured "
        "trace for UI display, and the final answer.")

    add_heading(doc, "3.1 Technology Stack", level=2)
    add_table(doc,
        headers=["Layer", "Choice", "Role"],
        rows=[
            ["Agent framework", "LangGraph (LangChain 0.3)", "Planner/executor graph + state"],
            ["LLM", "OpenAI GPT-4o-mini", "Planning, summarization, composition"],
            ["Embeddings", "OpenAI text-embedding-3-small", "FAISS vector store"],
            ["Vector store", "FAISS (local, on-disk)", "Patient H&P + Medline/WHO corpus"],
            ["Structured data", "SQLite (healthcare.db)", "Patients, doctors, schedule, logs"],
            ["UI", "Streamlit", "Chat, records, traces, evaluation"],
            ["Evaluation", "QAEvalChain", "LLM-as-judge grading of responses"],
        ],
        col_widths_in=[1.6, 2.6, 2.2],
    )

    add_heading(doc, "3.2 Data Layer", level=2)
    add_para(doc,
        "Structured data lives in a single SQLite file with tables for patients, doctors, schedule "
        "(half-hour slots across the next 14 weekdays), appointments, agent_runs, and "
        "patient_authorizations. The authorizations table maps each patient to one or more "
        "named contacts (with a relationship label such as daughter, son, husband, wife) who are "
        "permitted to act on their behalf. Unstructured data lives in two FAISS indices:")
    add_bullet(doc, "HP Index — chunks of the patient History & Physical PDFs, each chunk "
                    "metadata-tagged with patient_id so the agent can scope retrieval to one patient.")
    add_bullet(doc, "Medical Info Index — chunks of a seeded corpus of twelve markdown files "
                    "derived from Medline and WHO disease pages, each carrying source_url and "
                    "fetched_at frontmatter so citations remain traceable.")

    # ---- 4. Agent Design ----
    add_heading(doc, "4. Agent Design", level=1)

    add_heading(doc, "4.1 Tools", level=2)
    add_para(doc, "Eight tools are exposed to the agent. Tools marked PHI are gated by the "
                  "authorization layer described in section 5.")
    add_table(doc,
        headers=["Tool", "Purpose", "Backing store", "PHI"],
        rows=[
            ["lookup_patient",         "Fuzzy name / phone / description match", "SQLite + rapidfuzz", "no"],
            ["verify_caller",          "Caller-authorization check",   "SQLite",                "no"],
            ["get_patient_history",    "Retrieve + summarize H&P",     "FAISS (HP) + LLM",      "yes"],
            ["add_or_update_record",   "Create or modify a patient",   "SQLite",                "yes (update)"],
            ["list_doctor_slots",      "List open slots by specialty", "SQLite",                "no"],
            ["book_appointment",       "Book a specific slot",         "SQLite",                "yes"],
            ["cancel_appointment",     "Release a booked slot",        "SQLite",                "yes"],
            ["medical_info_search",    "RAG over Medline/WHO corpus",  "FAISS (medical) + LLM", "no*"],
        ],
        col_widths_in=[1.7, 2.3, 1.6, 0.6],
    )
    add_para(doc,
        "*medical_info_search is public, but its `condition` argument is gated: the executor "
        "refuses to research a topic the user did not name. This prevents the agent from leaking "
        "PHI by reading a patient's record, inferring their diagnosis, and surfacing it as "
        "general information. See section 5.4.",
        italic=True, size=10)

    add_heading(doc, "4.2 Planner", level=2)
    add_para(doc,
        "The planner is a single LLM call that produces a JSON array of ordered tool invocations. "
        "Each step names the tool, its arguments, and a one-line rationale. Arguments can reference "
        "the output of prior steps using the token $step{N}.path, resolved by the executor before "
        "each call. The planner sees the user query and any known patient context drawn from "
        "long-term memory.")

    add_heading(doc, "4.3 Executor", level=2)
    add_para(doc,
        "The executor iterates the plan, substitutes step references into arguments, calls the "
        "named tool, and appends the result — with success flag and latency — to step_results and "
        "to the structured trace. Tool errors are captured rather than raised, so the composer can "
        "still produce a useful reply that explains what went wrong. The executor also enforces "
        "the two authorization hard gates described in section 5: PHI-tool calls are blocked "
        "unless the patient_id has been verified earlier in the run, and medical_info_search "
        "calls are blocked when the requested condition cannot be derived from the user's own "
        "query. Blocked calls are recorded with `blocked_by_gate: True` and a structured reason, "
        "so policy enforcement is auditable rather than silent.")
    add_para(doc,
        "Step references support both dotted and bracket-indexed paths — `$step{1}.patient_id`, "
        "`$step{3}[0].slot_id`, `$step{2}.0.name` all resolve correctly — because the planner "
        "LLM naturally writes either style. Numeric arguments such as slot_id and patient_id are "
        "auto-coerced to int after substitution to avoid pydantic validation drift.")

    add_heading(doc, "4.4 Composer", level=2)
    add_para(doc,
        "The composer takes the plan and step results and produces one cohesive patient-facing "
        "reply: confirming any appointments booked (doctor name, date, time), preserving inline "
        "citations from the medical-info step, and flagging any tool failures with a suggested "
        "next action.")

    add_heading(doc, "4.5 Memory and Session State", level=2)
    add_para(doc,
        "Long-term memory is a rolling clinical summary on each patients row, updated by an LLM "
        "summarizer after each interaction, plus a queryable log of past agent runs. The result: "
        "on a follow-up query the planner sees the last few things the assistant did for that "
        "patient without re-reading every H&P chunk from scratch.")
    add_para(doc,
        "Session-level memory takes a deliberately lightweight form. Rather than feeding full "
        "conversation history into each planner call (which inflates token cost and risks the "
        "model fixating on irrelevant prior turns), the agent surfaces only two durable session "
        "facts: the set of patient_ids already verified for the current caller, and the "
        "patient_id currently being discussed. These are prepended to the planner prompt as "
        "plain-language directives. The Streamlit layer reads and writes these to "
        "st.session_state, and the executor's hard gate enforces the verification set "
        "independently — so even if the LLM ignored the directive, the gate would still block "
        "unauthorized PHI access.")
    add_code(doc,
        "# Build the planner's per-turn context block from session state.\n"
        "context_lines = []\n"
        "if caller := state.get('caller_name'):\n"
        "    context_lines.append(f'Caller: {caller}')\n"
        "if verified := state.get('verified_patient_ids'):\n"
        "    context_lines.append(\n"
        "        f'Already-verified patient_ids this session: {verified}. '\n"
        "        'Skip lookup_patient and verify_caller for these IDs.'\n"
        "    )\n"
        "if active_pid := state.get('patient_id'):\n"
        "    context_lines.append(\n"
        "        f'Active patient: patient_id={active_pid}. '\n"
        "        \"Use this when the user says 'him', 'her', 'this appointment', etc.\"\n"
        "    )\n"
        "context = '\\n'.join(context_lines)")

    # ---- 5. Authorization & PHI Protection ----
    add_heading(doc, "5. Authorization and PHI Protection", level=1)
    add_para(doc,
        "Healthcare conversations involve protected health information (PHI). Real clinic front "
        "desks verify callers before disclosing patient data — \"who am I speaking with, and "
        "what is your relationship to the patient?\" This system models that pattern with a "
        "dedicated authorization layer. Where it differs from a guardrail prompt is that the "
        "enforcement is in code, not in instructions. Even if the planner LLM is jailbroken or "
        "simply makes a poor choice, the executor's hard gate refuses to dispatch any "
        "PHI-touching tool call without a corresponding verification.")

    add_heading(doc, "5.1 Schema and seeded contacts", level=2)
    add_para(doc,
        "A new patient_authorizations table maps each patient to one or more named contacts, "
        "each with a relationship label and a phone number. The seed file populates a small "
        "sample for the demo: Mr. Harish Kumar's daughter Priya and son Anil; David Thompson's "
        "wife Sarah and son Michael; Anjali Mehra's husband Vikram; Ramesh Kulkarni's wife "
        "Sunita and son Anish. Verification matches on the contact name with a fuzzy "
        "(token-set) ratio threshold of 85, allowing for typos and middle initials.")

    add_heading(doc, "5.2 verify_caller tool", level=2)
    add_para(doc,
        "verify_caller is a tool the planner is required to call after lookup_patient and "
        "before any PHI-touching step. It returns a structured result containing `authorized` "
        "(bool), `relationship`, `contact_name`, the patient's name, and the office phone "
        "number to call for un-listed contacts. A successful verification adds the patient_id "
        "to the run's verified set; a failure returns authorized=false with a reason.")

    add_heading(doc, "5.3 Hard gate (verification)", level=2)
    add_para(doc,
        "The executor classifies each tool call. Calls to get_patient_history, "
        "book_appointment, and add_or_update_record (when a patient_id is supplied) are "
        "designated PHI tools. For any PHI tool, the executor checks whether the target "
        "patient_id is in the run's verified set; if not, it short-circuits the call and "
        "records a structured `blocked: true` result with reason `not_authorized`. The actual "
        "tool function never executes — no DB read, no FAISS query, no LLM call.")

    add_heading(doc, "5.4 Hard gate (inference leak)", level=2)
    add_para(doc,
        "A second, more subtle leak vector is topic inference. If the planner reads a patient's "
        "record (say, Anjali Mehra's H&P, which mentions a respiratory infection), and then "
        "decides to be helpful by searching public asthma information, the choice of topic "
        "itself reveals PHI — the agent has indirectly disclosed that Anjali likely has a "
        "respiratory issue, even while refusing to name her diagnosis. The system blocks this "
        "by gating medical_info_search on the requested condition: the executor refuses to "
        "research a condition that does not appear (loose substring + alias match) in the "
        "user's actual query. Twelve canonical aliases are recognized so abbreviations like "
        "CKD, T2DM, HTN, URI, COPD, CHF and others still authorize the corresponding full-name "
        "search.")
    add_code(doc,
        "# In the executor's loop:\n"
        "elif tool_name == 'medical_info_search':\n"
        "    cond = resolved_args.get('condition') or ''\n"
        "    user_q = state.get('user_query', '') or ''\n"
        "    if not _condition_derivable_from_query(cond, user_q):\n"
        "        block_out = {\n"
        "            'blocked': True,\n"
        "            'reason': 'condition_not_in_query',\n"
        "            'condition': cond,\n"
        "            'message': (\n"
        "                f\"Refused to research '{cond}' because the user did \"\n"
        "                'not name that condition. Researching a condition '\n"
        "                'inferred from a patient record would leak PHI.'\n"
        "            ),\n"
        "        }\n"
        "if block_out is not None:\n"
        "    # ...record the blocked step in step_results, then:\n"
        "    continue   # skip tool.invoke entirely")

    add_heading(doc, "5.5 Cross-turn persistence", level=2)
    add_para(doc,
        "Verification persists across follow-up turns within the same Streamlit session: a "
        "verified caregiver is not re-prompted between related questions. The Chat tab stores "
        "verified_patient_ids and last_patient_id in st.session_state and passes both into "
        "subsequent run_agent calls. Changing the caller name automatically wipes both — "
        "different person, no inherited trust. A 'New caller / clear session' button is "
        "exposed for explicit reset.")

    add_heading(doc, "5.6 Composer behavior under refusal", level=2)
    add_para(doc,
        "When PHI access is denied, the composer is instructed to apologize warmly, name the "
        "office phone number for unlisted-contact escalation, and never reveal any PHI from "
        "the blocked steps — including the patient's diagnosis, medications, or any inferred "
        "condition. Public medical information is still surfaced when the user themselves "
        "named the condition; otherwise the composer stops short of volunteering general "
        "information, since the topic itself would be a leak.")

    add_heading(doc, "5.7 Honest limitations of the auth layer", level=2)
    add_para(doc,
        "Name-only verification is demonstration-grade, not production-grade. Anyone who "
        "knows a relative's name could spoof this layer. A production deployment would "
        "require real identity proofing (one-time passwords to a verified phone, government "
        "ID match, single sign-on with a federated identity provider) and per-patient "
        "consent records that meet the relevant regulatory standard (HIPAA in the United "
        "States, the Digital Personal Data Protection Act in India). The architectural "
        "pattern — a code-level hard gate at the executor, with the LLM informed but not "
        "trusted as the sole enforcement point — is the durable contribution; the strength "
        "of the credential check is a separate knob to be turned up.")

    # ---- 6. LLMOps ----
    add_heading(doc, "6. LLMOps: Evaluation, Monitoring, and UI", level=1)

    add_heading(doc, "6.1 Evaluation harness", level=2)
    add_para(doc,
        "Evaluation runs a hand-labeled test set of 15 queries across the four capabilities "
        "(4 booking, 3 update, 4 history, 4 info), grades each agent response against a reference "
        "answer using QAEvalChain (LLM-as-judge), and aggregates accuracy per capability plus "
        "per-tool success rates and median / p95 latency. Results are serialized to a JSON file "
        "consumed by the Streamlit Evaluation tab and can be re-run any time from the terminal.")

    add_heading(doc, "6.2 Monitoring", level=2)
    add_para(doc,
        "Every agent run is persisted to the agent_runs table with its plan, trace, final answer, "
        "success flag, and total latency. The Agent Trace tab in the Streamlit UI renders any past "
        "run by id: the user query, the planner's JSON plan, each step with its arguments, output "
        "preview, latency, and success — the 'planning breakdown' the syllabus requires.")

    add_heading(doc, "6.3 Streamlit UI", level=2)
    add_para(doc, "Five tabs:")
    add_numbered(doc, "Chat — send a query, see the answer plus expandable plan and step results.")
    add_numbered(doc, "Patients — browse the patient table, create or update records.")
    add_numbered(doc, "Doctors & Appointments — doctor list, upcoming appointments, manual booking.")
    add_numbered(doc, "Agent Trace — past runs with plan and step-by-step breakdown.")
    add_numbered(doc, "Evaluation — accuracy by capability, per-tool success, latency percentiles.")

    # ---- 7. Flagship Scenario Walkthrough ----
    add_heading(doc, "7. Flagship Scenario Walkthrough", level=1)
    add_para(doc,
        'Given the caregiver Priya Kumar entering her name and the query "My 70-year-old '
        'father has chronic kidney disease. I want to book a nephrologist for him. Also, can '
        'you summarize latest treatment methods?" — the planner produces six steps:')
    for step in [
        "lookup_patient — resolve the description ('70 year old, chronic kidney disease') to "
        "Mr. Harish Kumar via fuzzy match against name, address, and clinical summary, with "
        "a +25 boost for exact age match.",
        "verify_caller — confirm Priya Kumar is on Mr. Kumar's authorized-contacts list "
        "(she is, as his daughter); the patient_id is added to the run's verified set.",
        "list_doctor_slots — return the next nephrologist slots in Bangalore (Mr. Kumar's "
        "city, derived from his address) with same-city preference and an out-of-city "
        "fallback flag.",
        "book_appointment — book the earliest slot for the verified patient.",
        "get_patient_history — pull H&P context for chronic kidney disease (allowed because "
        "Priya is verified).",
        "medical_info_search — retrieve and synthesize CKD treatment information; allowed "
        "because the user named CKD in their query.",
    ]:
        add_numbered(doc, step)
    add_para(doc,
        "The composer produces a single reply that opens with a warm verification "
        "acknowledgement (\"Thank you, Priya — you're verified as Mr. Kumar's daughter\"), "
        "confirms the booked appointment with doctor name, date, time, and city, lists up to "
        "five candidate slots so Priya can request a different one, and provides a cited "
        "summary of current CKD treatment options (ACE inhibitors / ARBs, SGLT2 inhibitors, "
        "finerenone, lifestyle measures, and renal replacement therapy in late stages). The "
        "Streamlit Appointments tab reflects the new booking immediately. On follow-up turns, "
        "verification carries forward — Priya is not re-prompted.")
    add_para(doc,
        "By contrast, the same query from an unauthorized caller (or no caller name at all) "
        "produces a different plan and outcome. lookup_patient may still find a candidate "
        "match, but verify_caller returns authorized=false; the executor then short-circuits "
        "every PHI step. Because the user did mention CKD by name, medical_info_search is "
        "still permitted, and the composer surfaces the public CKD information prefaced by "
        "the verification refusal and an instruction to call the office. If the unauthorized "
        "user had asked only \"tell me about Anjali\" with no condition named, the inference "
        "gate also blocks medical_info_search — preventing the agent from leaking that "
        "Anjali has a respiratory condition simply by researching asthma.")

    # ---- 8. Results ----
    add_heading(doc, "8. Results", level=1)
    add_para(doc,
        "Results below are representative placeholders that populate once the evaluation is run "
        "against your own OpenAI key. The Evaluation tab in the Streamlit app displays the live "
        "numbers and per-query breakdown after you click Run evaluation now.")
    add_table(doc,
        headers=["Capability", "# queries", "Target accuracy"],
        rows=[
            ["Appointment booking", "4", "100%"],
            ["Record update",       "3", ">= 90%"],
            ["History retrieval",   "4", ">= 85%"],
            ["Medical information", "4", ">= 85%"],
        ],
        col_widths_in=[2.4, 1.4, 2.2],
    )
    add_para(doc,
        "Tool-level success is also tracked (percentage of invocations that returned without an "
        "error). Median end-to-end latency is expected to fall between 3 and 8 seconds per turn "
        "for GPT-4o-mini.")

    # ---- 9. Limitations ----
    add_heading(doc, "9. Limitations", level=1)
    add_bullet(doc, "Caller verification is name-only and demonstration-grade (see section 5.7). "
                    "All patient data is synthetic and the prototype carries a visible disclaimer "
                    "that it is not clinical advice.")
    add_bullet(doc, "The medical-info corpus is static. Each response carries the fetched_at "
                    "date of its sources so freshness is explicit; a live Tavily / Serper / "
                    "Medline E-utilities fallback is a stretch goal.")
    add_bullet(doc, "Scheduling is mocked — the doctor calendar is seeded locally rather than "
                    "integrated with a real EHR or Google Calendar. Same-city preference and "
                    "out-of-city fallback are implemented; cross-clinic referrals are not.")
    add_bullet(doc, "Rebooking is supported via cancel_appointment (also a PHI-gated tool), "
                    "but the planner is not always reliable about chaining cancel-then-book in "
                    "one turn — the user may need to phrase the rebook explicitly (e.g. \"cancel "
                    "appointment 12 and book me into slot 1314 instead\").")
    add_bullet(doc, "Conversational memory is two state pointers (verified set + active "
                    "patient), not full chat history. This is deliberate to control token cost "
                    "and reduce drift, but it limits the agent's ability to refer back to "
                    "specific phrasing or numbers from earlier turns.")

    # ---- 10. Future Work ----
    add_heading(doc, "10. Future Work", level=1)
    add_bullet(doc, "Production-grade caller identity (one-time password to a verified phone, "
                    "government ID match, or federated SSO). The architectural pattern is in "
                    "place; the credential check is the knob to upgrade.")
    add_bullet(doc, "Robust single-turn rebooking: when the caregiver references a slot from "
                    "the prior turn, the planner should automatically chain "
                    "cancel_appointment(prior) + book_appointment(new) without the user having "
                    "to spell out both steps.")
    add_bullet(doc, "Live medical-info fallback via Tavily, Serper, or Medline E-utilities, "
                    "with a freshness check against the cached corpus.")
    add_bullet(doc, "Streaming token output in the Chat tab.")
    add_bullet(doc, "Prompt-injection hardening for patient-provided fields (free-text "
                    "summaries, address fields).")
    add_bullet(doc, "Fine-grained per-patient FAISS namespaces for scalable long-term memory.")
    add_bullet(doc, "Optional full chat-history memory for users who want it, with a clear "
                    "token-cost / drift tradeoff disclosed in the UI.")

    # ---- 11. Appendix ----
    add_heading(doc, "11. Appendix: Installation and Demo", level=1)
    add_code(doc,
        "python -m venv .venv && source .venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "cp .env.example .env   # add OPENAI_API_KEY\n"
        "python -m db.seed\n"
        "python -m src.embeddings --build\n"
        "streamlit run streamlit_app.py")
    add_para(doc,
        "All patient data is synthetic. Medical content is demo material, not clinical guidance.",
        italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p}")
